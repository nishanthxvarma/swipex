"""
SwipeX Layout-Aware Document & Text Extraction Engine.
Supports PDF (single & multi-column, tables, native text) and DOCX formats.
"""

import io
import re
import logging
from typing import Dict, List, Any, Optional, Tuple

logger = logging.getLogger(__name__)

class LayoutAwareExtractor:
    """
    Extracts structured, ordered text and layout metadata from raw PDF/DOCX bytes.
    """

    def extract_document(self, file_bytes: bytes, file_type: str, filename: str = "") -> Dict[str, Any]:
        """
        Main extraction entry point.
        Returns:
            {
                "raw_text": str,
                "pages": List[Dict],
                "page_count": int,
                "file_type": str,
                "is_scanned": bool,
                "confidence": float,
                "detected_columns": int,
                "tables_count": int
            }
        """
        ext = file_type.lower().replace(".", "")
        if ext == "pdf":
            return self._extract_pdf(file_bytes, filename)
        elif ext in ("docx", "doc"):
            return self._extract_docx(file_bytes, filename)
        else:
            return self._extract_plain_text(file_bytes)

    def _extract_pdf(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        pages_data = []
        full_text_chunks = []
        is_scanned = False
        total_chars = 0
        detected_columns = 1
        page_count = 0

        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            page_count = len(reader.pages)

            for p_idx, page in enumerate(reader.pages):
                # Try layout-mode extraction if available in pypdf
                try:
                    page_text = page.extract_text(extraction_mode="layout")
                except Exception:
                    page_text = page.extract_text() or ""

                if not page_text or len(page_text.strip()) < 20:
                    # Potential scanned page or sparse content
                    page_text = page_text or ""
                    
                cleaned_page_text = self._normalize_whitespace(page_text)
                total_chars += len(cleaned_page_text)

                # Heuristic column detection: presence of large central whitespace gap across lines
                lines = cleaned_page_text.split("\n")
                two_column_votes = 0
                for line in lines:
                    if re.search(r'\S\s{5,}\S', line):
                        two_column_votes += 1
                if two_column_votes > len(lines) * 0.25 and len(lines) > 10:
                    detected_columns = max(detected_columns, 2)

                pages_data.append({
                    "page_number": p_idx + 1,
                    "text": cleaned_page_text,
                    "char_count": len(cleaned_page_text),
                    "line_count": len(lines)
                })
                full_text_chunks.append(cleaned_page_text)

            if total_chars < 50 and page_count > 0:
                is_scanned = True
                confidence = 0.35
            else:
                confidence = 0.95 if total_chars > 300 else 0.70

        except Exception as e:
            logger.warning(f"pypdf extraction error for {filename}: {e}")
            fallback_text = file_bytes.decode("utf-8", errors="ignore")
            cleaned_fallback = self._normalize_whitespace(fallback_text)
            return {
                "raw_text": cleaned_fallback,
                "pages": [{"page_number": 1, "text": cleaned_fallback, "char_count": len(cleaned_fallback), "line_count": len(cleaned_fallback.splitlines())}],
                "page_count": 1,
                "file_type": "pdf",
                "is_scanned": len(cleaned_fallback) < 50,
                "confidence": 0.50,
                "detected_columns": 1,
                "tables_count": 0
            }

        # Deduplicate common headers/footers across pages if multiple pages
        if len(pages_data) > 1:
            full_text = self._deduplicate_headers_footers(pages_data)
        else:
            full_text = "\n\n".join(full_text_chunks)

        return {
            "raw_text": full_text.strip(),
            "pages": pages_data,
            "page_count": page_count,
            "file_type": "pdf",
            "is_scanned": is_scanned,
            "confidence": confidence,
            "detected_columns": detected_columns,
            "tables_count": 0
        }

    def _extract_docx(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        paragraphs_text = []
        tables_count = 0
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            
            # Extract paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs_text.append(p.text.strip())

            # Extract table cells preserving text
            for table in doc.tables:
                tables_count += 1
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        # Join distinct non-empty cells
                        unique_cells = []
                        for c in row_cells:
                            if c not in unique_cells:
                                unique_cells.append(c)
                        paragraphs_text.append(" | ".join(unique_cells))

            full_text = "\n".join(paragraphs_text)
            total_chars = len(full_text)
            confidence = 0.98 if total_chars > 200 else 0.75

            return {
                "raw_text": self._normalize_whitespace(full_text),
                "pages": [{"page_number": 1, "text": full_text, "char_count": total_chars, "line_count": len(paragraphs_text)}],
                "page_count": 1,
                "file_type": "docx",
                "is_scanned": False,
                "confidence": confidence,
                "detected_columns": 1,
                "tables_count": tables_count
            }
        except Exception as e:
            logger.warning(f"python-docx extraction error for {filename}: {e}")
            fallback_text = file_bytes.decode("utf-8", errors="ignore")
            cleaned = self._normalize_whitespace(fallback_text)
            return {
                "raw_text": cleaned,
                "pages": [{"page_number": 1, "text": cleaned, "char_count": len(cleaned), "line_count": len(cleaned.splitlines())}],
                "page_count": 1,
                "file_type": "docx",
                "is_scanned": False,
                "confidence": 0.50,
                "detected_columns": 1,
                "tables_count": 0
            }

    def _extract_plain_text(self, file_bytes: bytes) -> Dict[str, Any]:
        text = file_bytes.decode("utf-8", errors="ignore")
        cleaned = self._normalize_whitespace(text)
        return {
            "raw_text": cleaned,
            "pages": [{"page_number": 1, "text": cleaned, "char_count": len(cleaned), "line_count": len(cleaned.splitlines())}],
            "page_count": 1,
            "file_type": "txt",
            "is_scanned": False,
            "confidence": 0.90,
            "detected_columns": 1,
            "tables_count": 0
        }

    def _normalize_whitespace(self, text: str) -> str:
        """
        Normalizes unicode spaces, trims trailing whitespaces, and preserves clean newlines.
        """
        if not text:
            return ""
        # Replace non-breaking spaces
        text = text.replace('\xa0', ' ').replace('\u200b', '')
        # Replace multiple spaces within lines while preserving line structure
        lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in text.splitlines()]
        # Remove consecutive blank lines (more than 2)
        result_lines = []
        blank_count = 0
        for line in lines:
            if not line:
                blank_count += 1
                if blank_count <= 2:
                    result_lines.append("")
            else:
                blank_count = 0
                result_lines.append(line)
        return "\n".join(result_lines)

    def _deduplicate_headers_footers(self, pages: List[Dict]) -> str:
        """
        Detects and removes repeating header and footer lines across multi-page documents.
        """
        if len(pages) <= 1:
            return "\n\n".join(p["text"] for p in pages)

        # Check first and last 2 lines of each page
        first_lines = [p["text"].split("\n")[0] for p in pages if p["text"].split("\n")]
        last_lines = [p["text"].split("\n")[-1] for p in pages if p["text"].split("\n")]

        common_header = first_lines[0] if len(set(first_lines)) == 1 and len(first_lines[0]) < 80 else None
        common_footer = last_lines[0] if len(set(last_lines)) == 1 and len(last_lines[0]) < 80 else None

        cleaned_pages = []
        for p in pages:
            lines = p["text"].split("\n")
            if common_header and lines and lines[0] == common_header:
                lines = lines[1:]
            if common_footer and lines and lines[-1] == common_footer:
                lines = lines[:-1]
            cleaned_pages.append("\n".join(lines))

        return "\n\n".join(cleaned_pages)

layout_extractor = LayoutAwareExtractor()
